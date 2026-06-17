from __future__ import annotations

import hashlib
import json
import os
import re
from pathlib import Path
from typing import Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config import DATA_DIR, RECOMMENDATIONS_DIR
from skill_analysis.analysis import extract_job_skills, extract_resume_skills, job_row_from_rec
from tailoring_resume.gemini_utils import call_gemini_json, gemini_is_configured

CACHE_FILE = DATA_DIR / "processed" / "tailored_resume_cache.json"
OUTPUT_DIR = RECOMMENDATIONS_DIR / "tailored_resumes"
DEFAULT_MODEL = os.environ.get("TAILORED_RESUME_MODEL", "gemini-2.5-flash")
PROMPT_VERSION = "tailored-resume-v1"

TAILORING_PROMPT = """
You are an expert ATS resume strategist.

Your task is to tailor the user's existing resume to a target job without inventing facts.

Hard rules:
- Do not add skills, tools, certifications, projects, employers, education, metrics, or achievements that are not supported by the original resume input.
- Do not claim the user used a technology unless it appears in the original resume input.
- Do not create fake numbers, dates, experience, or responsibilities.
- You may reorder sections, tighten wording, rewrite bullets for ATS clarity, surface relevant existing skills, and de-emphasize irrelevant content.
- Keep the resume realistic, concise, recruiter-friendly, and faithful to the source material.
- If a detail is not supported, omit it instead of guessing.

Output requirements:
- Return ONLY valid JSON.
- `resume_markdown` must contain the full tailored resume in markdown.
- Use this structure when possible:
  # Full Name
  Contact line
  ## Professional Summary
  paragraph
  ## Technical Skills
  - item
  ## Experience
  ### Role / Organization
  - bullet
  ## Projects
  ### Project Name
  - bullet
  ## Education
  paragraph
- Omit unsupported sections instead of inventing them.
""".strip()


def tailored_resume_is_configured() -> bool:
    return gemini_is_configured()


def _load_cache() -> Dict:
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return data
        except Exception:
            return {}
    return {}


def _save_cache(cache: Dict) -> None:
    CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2, ensure_ascii=False)


def _hash_payload(payload: Dict, model: str) -> str:
    raw = json.dumps(
        {"version": PROMPT_VERSION, "model": model, "payload": payload},
        sort_keys=True,
        ensure_ascii=False,
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _slugify(text: str, default: str = "resume") -> str:
    cleaned = re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")
    return cleaned[:80] or default


def _responses_schema() -> Dict:
    return {
        "type": "json_schema",
        "name": "tailored_resume",
        "strict": True,
        "schema": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "resume_markdown": {"type": "string"},
                "ats_focus": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "tailoring_summary": {"type": "string"},
            },
            "required": ["resume_markdown", "ats_focus", "tailoring_summary"],
        },
    }


def _build_payload(preprocessed: dict, rec: dict) -> Dict:
    job_row = job_row_from_rec(rec)
    full_description = str(job_row.get("Job Description") or rec.get("description") or "").strip()
    job_skills = extract_job_skills({**rec, "description": full_description})

    sections = {
        key: value
        for key, value in (preprocessed.get("sections") or {}).items()
        if isinstance(value, str) and value.strip()
    }

    return {
        "resume": {
            "filename": preprocessed.get("filename", ""),
            "raw_text": str(preprocessed.get("raw_text", ""))[:12000],
            "sections": sections,
            "current_skills": extract_resume_skills(preprocessed),
        },
        "job": {
            "title": rec.get("title", ""),
            "company": rec.get("company", ""),
            "domain": rec.get("domain", ""),
            "job_description": full_description[:8000],
            "matching_skills": rec.get("matching_skills") or [],
            "missing_skills": rec.get("missing_skills") or [],
            "job_skills": job_skills,
        },
    }


def _call_tailor_llm(payload: Dict, model: str) -> Dict:
    prompt = f"""
{TAILORING_PROMPT}

Input JSON:
{json.dumps(payload, ensure_ascii=False, indent=2)}

Return ONLY valid JSON in this format:
{json.dumps(_responses_schema()["schema"], indent=2)}
"""
    return call_gemini_json(prompt, model, retries=1)


def _output_path(preprocessed: dict, rec: dict, cache_key: str) -> Path:
    resume_stem = Path(preprocessed.get("filename", "resume")).stem
    job_slug = _slugify(f"{rec.get('title', '')}-{rec.get('company', '')}", default="job")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR / f"{_slugify(resume_stem)}__{job_slug}__{cache_key[:10]}.docx"


def _set_base_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def _write_markdown_to_docx(markdown: str, output_path: Path) -> Path:
    doc = Document()
    _set_base_styles(doc)

    header_mode = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            header_mode = True
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            header_mode = False
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            header_mode = False
            continue

        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        paragraph = doc.add_paragraph()
        if header_mode:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            run.font.size = Pt(10)
        else:
            paragraph.add_run(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_tailored_resume(preprocessed: dict, rec: dict, model: str = DEFAULT_MODEL) -> Dict:
    if not tailored_resume_is_configured():
        raise RuntimeError("Tailored resume generation requires GEMINI_API_KEY.")

    payload = _build_payload(preprocessed, rec)
    cache_key = _hash_payload(payload, model)
    cache = _load_cache()

    cached = cache.get(cache_key)
    if isinstance(cached, dict):
        cached_docx_path = cached.get("docx_path") or ""
        output_path = Path(cached_docx_path) if cached_docx_path else None
        markdown = cached.get("resume_markdown", "")
        if markdown and (output_path is None or not output_path.exists()):
            output_path = _output_path(preprocessed, rec, cache_key)
            _write_markdown_to_docx(markdown, output_path)
            cached["docx_path"] = str(output_path)
            cache[cache_key] = cached
            _save_cache(cache)
        return cached

    result = _call_tailor_llm(payload, model)
    output_path = _output_path(preprocessed, rec, cache_key)
    _write_markdown_to_docx(result["resume_markdown"], output_path)

    artifact = {
        **result,
        "used_llm": True,
        "model": result.get("_llm_model") or model,
        "backend": result.get("_llm_backend") or "gemini",
        "docx_path": str(output_path),
        "target_title": rec.get("title", ""),
        "target_company": rec.get("company", ""),
        "cache_key": cache_key,
    }
    cache[cache_key] = artifact
    _save_cache(cache)
    return artifact
